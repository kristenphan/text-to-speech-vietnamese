import PyPDF2
import sys
import requests
from googleapiclient.discovery import build
from docx import Document


# 1. Extract text from the PDF by chapters and translate to Vietnamese
google_translate_service = build(
    "translate",
    "v2",
    developerKey="{api-key}")

pages: list[str] = []

with open("./books/David-A-Phillips-The-Complete-Book-of-Numerology.pdf", "rb") as f:
    reader = PyPDF2.PdfReader(f)
    current_chapter_num = None

    print(f'Book metadata: {reader.metadata}')
    print(f'# pages: {len(reader.pages)}')

    for page_num in range(5, len(reader.pages) - 6):  # Skip everything before "Tables of Contents" and after "About Author"
        print(f"Reading page {page_num}...")
        page_content = reader.pages[page_num].extract_text()
 
        page_translation_response = google_translate_service.translations()\
                .list(source="en", target="vi", q=[page_content])\
                .execute()
        translated_page = page_translation_response['translations'][0]['translatedText']

        pages.append(translated_page)

# 2. Write the translated text to .docx
doc = Document()

doc.add_heading("Translated Chapters", level=1)

for page in pages:
    doc.add_paragraph(page)
    doc.add_page_break()

doc.save("./out/translated.docx")

# 3. Generate Vietnamese audio files for each chapter
# Need to build and run Docker container: follow ./vits_tts_vietnamese/README.md
for (page_num, page_content) in enumerate(pages):
    url = "http://localhost:5004/tts"
    params = {
        "text": page_content,
        "speed": "slow"
    }
    response = requests.get(
        url=url,
        params=params)

    if response.status_code == 200:
        audio_url = response.json().get("audio_url")
        with open(f"./out/Trang_{page_num}.wav", "wb") as f:
            f.write(requests.get(audio_url).content)
            print("Audio file saved successfully.")
    else:
        print(f"Failed to generate audio: {response.status_code}")